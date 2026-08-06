from datetime import datetime
from pathlib import Path


class ExportService:
    def export_txt(
        self,
        destination: Path,
        text: str,
    ) -> None:
        destination.parent.mkdir(
            parents=True,
            exist_ok=True,
        )
        destination.write_text(
            text,
            encoding="utf-8-sig",
        )

    def export_docx(
        self,
        destination: Path,
        text: str,
        source_name: str = "",
    ) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        destination.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(55)
        section.right_margin = Pt(55)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(
            "AUDITOR IA - TRANSCRIPTOR"
        )
        title_run.bold = True
        title_run.font.size = Pt(16)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(
            "TRANSCRIPCIÓN DE AUDIO"
        )
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(12)

        if source_name:
            source = document.add_paragraph()
            source.add_run("ARCHIVO: ").bold = True
            source.add_run(source_name)

        generated = document.add_paragraph()
        generated.add_run(
            "FECHA DE EXPORTACIÓN: "
        ).bold = True
        generated.add_run(
            datetime.now().strftime(
                "%d-%m-%Y %H:%M"
            )
        )

        document.add_paragraph()

        for line in text.splitlines():
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.size = Pt(10.5)

        document.save(destination)
